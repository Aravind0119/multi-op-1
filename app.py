# import os
# import json
# import pandas as pd
# import base64
# import xml.etree.ElementTree as ET
# from flask import Flask, request, jsonify, render_template_string
# from docx import Document
# from PyPDF2 import PdfReader

# app = Flask(__name__)

# stored_data = None


# # -----------------------------
# # Extract PDF
# # -----------------------------
# def extract_pdf_text(file):
#     reader = PdfReader(file)
#     text = ""
#     for page in reader.pages:
#         text += page.extract_text() or ""
#     return {"pdf_text": text}


# # -----------------------------
# # Extract DOCX
# # -----------------------------
# def extract_docx_text(file):
#     doc = Document(file)
#     text = "\n".join([p.text for p in doc.paragraphs])
#     return {"docx_text": text}


# # -----------------------------
# # XML to dict
# # -----------------------------
# def xml_to_dict(element):
#     return {
#         element.tag: {
#             "text": element.text,
#             "attributes": element.attrib,
#             "children": [xml_to_dict(child) for child in element]
#         }
#     }


# # -----------------------------
# # Professional Home Page
# # -----------------------------
# @app.route("/")
# def home():
#     return render_template_string("""
#     <!DOCTYPE html>
#     <html>
#     <head>
#         <title>Webhook Dashboard</title>
#         <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css" rel="stylesheet">
#     </head>
#     <body class="bg-light">

#     <div class="container mt-5">
#         <div class="card shadow-lg p-4">
#             <h2 class="text-center mb-4">📂 Webhook File Dashboard</h2>

#             <div class="text-center">
#                 <a href="/upload" class="btn btn-primary btn-lg m-2">Upload File</a>
#                 <a href="/webhook" class="btn btn-success btn-lg m-2">View Stored JSON</a>
#             </div>
#         </div>
#     </div>

#     </body>
#     </html>
#     """)


# # -----------------------------
# # Upload Page
# # -----------------------------
# @app.route("/upload", methods=["GET", "POST"])
# def upload_file():
#     global stored_data

#     if request.method == "GET":
#         return render_template_string("""
#         <!DOCTYPE html>
#         <html>
#         <head>
#             <title>Upload File</title>
#             <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css" rel="stylesheet">
#             <style>
#                 .drop-zone {
#                     border: 3px dashed #0d6efd;
#                     border-radius: 15px;
#                     padding: 50px;
#                     text-align: center;
#                     cursor: pointer;
#                     transition: 0.3s;
#                     background-color: #f8f9fa;
#                 }
#                 .drop-zone.dragover {
#                     background-color: #e7f1ff;
#                 }
#                 .file-name {
#                     margin-top: 15px;
#                     font-weight: bold;
#                     color: #198754;
#                 }
#             </style>
#         </head>
#         <body class="bg-light">

#         <div class="container mt-5">
#             <div class="card shadow-lg p-4">
#                 <h3 class="text-center mb-4">📂 Drag & Drop File Upload</h3>

#                 <form method="POST" enctype="multipart/form-data" id="uploadForm">
#                     <div class="drop-zone" id="dropZone">
#                         <p>Drag & Drop your file here</p>
#                         <p>or</p>
#                         <button type="button" class="btn btn-primary" onclick="document.getElementById('fileInput').click()">Browse File</button>
#                         <input type="file" name="file" id="fileInput" hidden required>
#                         <div class="file-name" id="fileName"></div>
#                     </div>

#                     <div class="text-center mt-4">
#                         <button type="submit" class="btn btn-success">Upload & Convert</button>
#                         <a href="/" class="btn btn-secondary">Back</a>
#                     </div>
#                 </form>
#             </div>
#         </div>

#         <script>
#             const dropZone = document.getElementById("dropZone");
#             const fileInput = document.getElementById("fileInput");
#             const fileName = document.getElementById("fileName");

#             dropZone.addEventListener("dragover", (e) => {
#                 e.preventDefault();
#                 dropZone.classList.add("dragover");
#             });

#             dropZone.addEventListener("dragleave", () => {
#                 dropZone.classList.remove("dragover");
#             });

#             dropZone.addEventListener("drop", (e) => {
#                 e.preventDefault();
#                 dropZone.classList.remove("dragover");

#                 const files = e.dataTransfer.files;
#                 fileInput.files = files;
#                 fileName.textContent = "Selected File: " + files[0].name;
#             });

#             fileInput.addEventListener("change", () => {
#                 fileName.textContent = "Selected File: " + fileInput.files[0].name;
#             });
#         </script>

#         </body>
#         </html>
#         """)

#     # -----------------------------
#     # POST Logic (Keep Same)
#     # -----------------------------
#     file = request.files["file"]
#     filename = file.filename.lower()

#     try:
#         if filename.endswith(".csv"):
#             df = pd.read_csv(file)
#             stored_data = df.to_dict(orient="records")

#         elif filename.endswith(".json"):
#             stored_data = json.load(file)

#         elif filename.endswith(".xlsx"):
#             df = pd.read_excel(file)
#             stored_data = df.to_dict(orient="records")

#         elif filename.endswith(".txt"):
#             stored_data = {"text_content": file.read().decode("utf-8", errors="ignore")}

#         elif filename.endswith(".xml"):
#             tree = ET.parse(file)
#             root = tree.getroot()
#             stored_data = xml_to_dict(root)

#         elif filename.endswith(".pdf"):
#             stored_data = extract_pdf_text(file)

#         elif filename.endswith(".docx"):
#             stored_data = extract_docx_text(file)

#         else:
#             stored_data = {
#                 "filename": filename,
#                 "base64_data": base64.b64encode(file.read()).decode("utf-8")
#             }

#         return render_template_string("""
#         <div style="text-align:center;margin-top:50px;">
#             <h3 style="color:green;">✅ File Converted Successfully</h3>
#             <a href="/webhook" style="font-size:18px;">View JSON</a><br><br>
#             <a href="/">Back to Dashboard</a>
#         </div>
#         """)

#     except Exception as e:
#         return f"Error: {str(e)}"

# # -----------------------------
# # Pretty JSON Viewer
# # -----------------------------
# @app.route("/webhook")
# def webhook():
#     global stored_data

#     if stored_data is None:
#         return jsonify({"message": "No data stored yet"})

#     formatted_json = json.dumps(stored_data, indent=4)

#     return render_template_string("""
#     <!DOCTYPE html>
#     <html>
#     <head>
#         <title>Stored JSON</title>
#         <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css" rel="stylesheet">
#         <style>
#             pre {
#                 background-color: #1e1e1e;
#                 color: #00ff99;
#                 padding: 20px;
#                 border-radius: 8px;
#                 overflow-x: auto;
#             }
#         </style>
#     </head>
#     <body class="bg-light">

#     <div class="container mt-5">
#         <div class="card shadow-lg p-4">
#             <h3 class="mb-3">📄 Stored JSON Output</h3>
#             <pre>{{ data }}</pre>
#             <a href="/" class="btn btn-secondary mt-3">Back</a>
#         </div>
#     </div>

#     </body>
#     </html>
#     """, data=formatted_json)


# # -----------------------------
# # Run for Render
# # -----------------------------
# if __name__ == "__main__":
#     port = int(os.environ.get("PORT", 5000))
#     app.run(host="0.0.0.0", port=port)

# import os
# import json
# import pandas as pd
# import base64
# import xml.etree.ElementTree as ET
# from flask import Flask, request, jsonify, render_template_string
# from docx import Document
# from PyPDF2 import PdfReader

# app = Flask(__name__)

# stored_data = {}

# # -----------------------------
# # File Converters
# # -----------------------------

# def extract_pdf_text(file):
#     reader = PdfReader(file)
#     text = ""
#     for page in reader.pages:
#         text += page.extract_text() or ""
#     return {"pdf_text": text}

# def extract_docx_text(file):
#     doc = Document(file)
#     return {"docx_text": "\n".join([p.text for p in doc.paragraphs])}

# def xml_to_dict(element):
#     return {
#         element.tag: {
#             "text": element.text,
#             "attributes": element.attrib,
#             "children": [xml_to_dict(child) for child in element]
#         }
#     }

# # -----------------------------
# # PROFESSIONAL DASHBOARD
# # -----------------------------

# @app.route("/")
# def home():
#     return render_template_string("""
# <!DOCTYPE html>
# <html>
# <head>
# <title>Professional Webhook Dashboard</title>
# <meta name="viewport" content="width=device-width, initial-scale=1">
# <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css" rel="stylesheet">
# <style>
# body { background: linear-gradient(135deg, #1f4037, #99f2c8); min-height:100vh; }
# .card { border-radius: 15px; }
# .drop-zone {
#     border: 3px dashed #0d6efd;
#     border-radius: 15px;
#     padding: 50px;
#     text-align: center;
#     background: #f8f9fa;
#     transition: 0.3s;
# }
# .drop-zone.dragover {
#     background: #e7f1ff;
# }
# .file-list {
#     margin-top: 15px;
# }
# pre {
#     background: #111;
#     color: #00ff9d;
#     padding: 20px;
#     border-radius: 10px;
#     max-height: 400px;
#     overflow:auto;
# }
# </style>
# </head>
# <body>

# <div class="container py-5">
# <div class="card shadow-lg p-4">

# <h2 class="text-center mb-4">📂 Multi File Webhook Dashboard</h2>

# <form method="POST" action="/upload" enctype="multipart/form-data" id="uploadForm">

# <div class="drop-zone" id="dropZone">
#     <h5>Drag & Drop Files Here</h5>
#     <p>or</p>
#     <button type="button" class="btn btn-primary" onclick="fileInput.click()">Browse Files</button>
#     <input type="file" id="fileInput" name="file" multiple hidden required>
#     <div class="file-list" id="fileList"></div>
# </div>

# <div class="text-center mt-4">
#     <button type="submit" class="btn btn-success">Upload & Convert</button>
#     <a href="/webhook" class="btn btn-dark">View JSON</a>
# </div>

# </form>

# </div>
# </div>

# <script>
# const dropZone = document.getElementById("dropZone");
# const fileInput = document.getElementById("fileInput");
# const fileList = document.getElementById("fileList");

# dropZone.addEventListener("dragover", e => {
#     e.preventDefault();
#     dropZone.classList.add("dragover");
# });

# dropZone.addEventListener("dragleave", () => {
#     dropZone.classList.remove("dragover");
# });

# dropZone.addEventListener("drop", e => {
#     e.preventDefault();
#     dropZone.classList.remove("dragover");
#     fileInput.files = e.dataTransfer.files;
#     updateFileList();
# });

# fileInput.addEventListener("change", updateFileList);

# function updateFileList() {
#     fileList.innerHTML = "";
#     Array.from(fileInput.files).forEach(file => {
#         fileList.innerHTML += "<div>📄 " + file.name + "</div>";
#     });
# }
# </script>

# </body>
# </html>
# """)

# # -----------------------------
# # Upload Logic
# # -----------------------------

# @app.route("/upload", methods=["POST"])
# def upload():
#     global stored_data

#     # 🔴 CLEAR PREVIOUS DATA
#     stored_data = {}

#     files = request.files.getlist("file")

#     for file in files:
#         filename = file.filename.lower()

#         try:
#             if filename.endswith(".csv"):
#                 df = pd.read_csv(file)
#                 stored_data[filename] = df.to_dict(orient="records")

#             elif filename.endswith(".json"):
#                 stored_data[filename] = json.load(file)

#             elif filename.endswith(".xlsx"):
#                 df = pd.read_excel(file)
#                 stored_data[filename] = df.to_dict(orient="records")

#             elif filename.endswith(".txt"):
#                 stored_data[filename] = {
#                     "text": file.read().decode("utf-8", errors="ignore")
#                 }

#             elif filename.endswith(".xml"):
#                 tree = ET.parse(file)
#                 stored_data[filename] = xml_to_dict(tree.getroot())

#             elif filename.endswith(".pdf"):
#                 stored_data[filename] = extract_pdf_text(file)

#             elif filename.endswith(".docx"):
#                 stored_data[filename] = extract_docx_text(file)

#             else:
#                 stored_data[filename] = {
#                     "base64_data": base64.b64encode(file.read()).decode("utf-8")
#                 }

#         except Exception as e:
#             stored_data[filename] = {"error": str(e)}

#     return """
#     <div style='text-align:center;margin-top:50px;'>
#     <h3 style='color:green;'>✅ Files Converted Successfully</h3>
#     <a href='/'>Back</a> | <a href='/webhook'>View JSON</a>
#     </div>
#     """


# # -----------------------------
# # JSON Viewer
# # -----------------------------

# @app.route("/webhook")
# def webhook():
#     if not stored_data:
#         return "<h3 style='text-align:center;margin-top:50px;'>No files uploaded yet.</h3>"

#     formatted = json.dumps(stored_data, indent=4)

#     return f"""
#     <div style='padding:40px;background:#111;color:#00ff9d;font-family:monospace;'>
#     <h2>Stored JSON Output</h2>
#     <pre>{formatted}</pre>
#     <a href='/' style='color:white;'>Back</a>
#     </div>
#     """

# # -----------------------------
# # Run
# # -----------------------------

# if __name__ == "__main__":
#     port = int(os.environ.get("PORT", 5000))
#     app.run(host="0.0.0.0", port=port)

# from flask import Flask, request, render_template_string, jsonify
# import re
# from datetime import datetime

# app = Flask(__name__)

# stored_data = {}

# # ------------------- BANK KYC FORM -------------------

# HTML_FORM = """
# <!DOCTYPE html>
# <html>
# <head>
# <title>Bank Customer Onboarding - KYC</title>
# <style>
# body { font-family: Arial; background:#f4f6f9; }
# .container { width:80%; margin:auto; background:white; padding:30px; margin-top:20px; }
# h2 { background:#002b5c; color:white; padding:10px; }
# h3 { border-bottom:1px solid #ccc; padding-top:15px; }
# input, select { width:100%; padding:8px; margin-bottom:10px; }
# button { padding:12px 25px; background:#003399; color:white; border:none; cursor:pointer; }
# .section { margin-bottom:25px; }
# </style>
# </head>
# <body>

# <div class="container">
# <h2>Customer Due Diligence (CDD) / KYC Application</h2>

# <form method="POST">

# <div class="section">
# <h3>1. Identification Details</h3>

# Customer ID:
# <input type="text" name="customerId" required>

# Full Name:
# <input type="text" name="fullName" required>

# Father / Spouse Name:
# <input type="text" name="guardianName" required>

# Date of Birth:
# <input type="date" name="dob" required>

# Gender:
# <select name="gender" required>
# <option value="">Select</option>
# <option>Male</option>
# <option>Female</option>
# <option>Other</option>
# </select>

# Nationality:
# <input type="text" name="nationality" required>

# PAN Number:
# <input type="text" name="panNo" required>

# Aadhaar Number:
# <input type="text" name="aadhaarNo" required>
# </div>

# <div class="section">
# <h3>2. Contact Details</h3>

# Mobile Number:
# <input type="text" name="mobile" required>

# Email:
# <input type="email" name="email" required>

# Residential Address:
# <input type="text" name="address" required>

# City:
# <input type="text" name="city" required>

# State:
# <input type="text" name="state" required>

# PIN Code:
# <input type="text" name="pincode" required>
# </div>

# <div class="section">
# <h3>3. Employment & Financial Details</h3>

# Occupation:
# <select name="occupation" required>
# <option value="">Select</option>
# <option>Salaried</option>
# <option>Self Employed</option>
# <option>Business</option>
# <option>Professional</option>
# <option>Student</option>
# <option>Retired</option>
# </select>

# Employer Name:
# <input type="text" name="employer">

# Annual Income (INR):
# <input type="number" name="annualIncome" required>

# Source of Funds:
# <input type="text" name="sourceOfFunds" required>
# </div>

# <div class="section">
# <h3>4. Regulatory Declarations</h3>

# Politically Exposed Person (PEP):
# <select name="pepStatus" required>
# <option value="">Select</option>
# <option>No</option>
# <option>Yes</option>
# </select>

# Tax Residency Country:
# <input type="text" name="taxCountry" required>

# <input type="checkbox" name="fatcaDeclaration" required>
# I confirm FATCA/CRS compliance.

# <br><br>
# <input type="checkbox" name="customerDeclaration" required>
# I confirm information provided is true.

# Signature:
# <input type="text" name="signature" required>

# Date:
# <input type="date" name="signatureDate" required>
# </div>

# <button type="submit">Submit Application</button>

# </form>
# </div>

# </body>
# </html>
# """

# from flask import Flask, request, render_template_string, jsonify
# import re
# from datetime import datetime
# import uuid

# app = Flask(__name__)

# # ---------------- MASKING UTILITIES ----------------

# def mask_aadhaar(aadhaar):
#     if aadhaar and len(aadhaar) == 12:
#         return "XXXX XXXX " + aadhaar[-4:]
#     return None

# def mask_pan(pan):
#     if pan and len(pan) == 10:
#         return pan[:5] + "XXXX"
#     return None

# def mask_mobile(mobile):
#     if mobile and len(mobile) == 10:
#         return "XXXXXX" + mobile[-4:]
#     return None

# # ---------------- RBI-STYLE PROFESSIONAL KYC FORM ----------------

# HTML_FORM = """
# <!DOCTYPE html>
# <html>
# <head>
# <title>RBI Compliant Customer KYC Form</title>
# <style>
# body { font-family: Arial; background:#eef2f7; }
# .container { width:75%; margin:auto; background:white; padding:30px; margin-top:20px; border-radius:6px; box-shadow:0 0 10px #ccc; }
# h1 { text-align:center; color:#002b5c; }
# h3 { border-bottom:2px solid #002b5c; padding-top:20px; color:#002b5c; }
# label { font-weight:bold; }
# input, select { width:100%; padding:8px; margin:6px 0 12px 0; }
# button { padding:12px 25px; background:#002b5c; color:white; border:none; cursor:pointer; }
# </style>
# </head>
# <body>

# <div class="container">
# <h1>Customer Due Diligence (CDD) – KYC Application</h1>

# <form method="POST">

# <h3>1. Customer Identification Details</h3>
# <label>Customer ID</label>
# <input type="text" name="customerId" required>

# <label>Full Name</label>
# <input type="text" name="fullName" required>

# <label>Father / Spouse Name</label>
# <input type="text" name="guardianName" required>

# <label>Date of Birth</label>
# <input type="date" name="dob" required>

# <label>Gender</label>
# <select name="gender" required>
# <option value="">Select</option>
# <option>Male</option>
# <option>Female</option>
# <option>Other</option>
# </select>

# <label>Nationality</label>
# <input type="text" name="nationality" required>

# <h3>2. Officially Valid Documents (OVD)</h3>
# <label>PAN Number</label>
# <input type="text" name="panNo" required>

# <label>Aadhaar Number</label>
# <input type="text" name="aadhaarNo" required>

# <h3>3. Contact Information</h3>
# <label>Mobile Number</label>
# <input type="text" name="mobile" required>

# <label>Email Address</label>
# <input type="email" name="email" required>

# <label>Residential Address</label>
# <input type="text" name="address" required>

# <label>City</label>
# <input type="text" name="city" required>

# <label>State</label>
# <input type="text" name="state" required>

# <label>PIN Code</label>
# <input type="text" name="pincode" required>

# <h3>4. Occupation & Financial Profile</h3>
# <label>Occupation</label>
# <select name="occupation" required>
# <option value="">Select</option>
# <option>Salaried</option>
# <option>Self Employed</option>
# <option>Business</option>
# <option>Professional</option>
# <option>Student</option>
# <option>Retired</option>
# </select>

# <label>Employer Name</label>
# <input type="text" name="employer">

# <label>Annual Income (INR)</label>
# <input type="number" name="annualIncome" required>

# <label>Source of Funds</label>
# <input type="text" name="sourceOfFunds" required>

# <h3>5. Regulatory Declarations</h3>
# <label>Politically Exposed Person (PEP)</label>
# <select name="pepStatus" required>
# <option value="">Select</option>
# <option>No</option>
# <option>Yes</option>
# </select>

# <label>Tax Residency Country</label>
# <input type="text" name="taxCountry" required>

# <label>Declaration</label>
# <input type="checkbox" required> I confirm that the information provided is true and correct.

# <br><br>
# <button type="submit">Submit KYC Application</button>

# </form>
# </div>
# </body>
# </html>
# """

# # ---------------- ROUTE ----------------

# @app.route("/", methods=["GET", "POST"])
# def kyc():

#     if request.method == "GET":
#         return render_template_string(HTML_FORM)

#     data = request.form.to_dict()
#     errors = []

#     # -------- RBI Validation Checks --------

#     if not re.fullmatch(r"\d{12}", data.get("aadhaarNo", "")):
#         errors.append("Invalid Aadhaar Number")

#     if not re.fullmatch(r"[A-Z]{5}[0-9]{4}[A-Z]", data.get("panNo", "")):
#         errors.append("Invalid PAN Number")

#     if not re.fullmatch(r"\d{10}", data.get("mobile", "")):
#         errors.append("Invalid Mobile Number")

#     if not re.fullmatch(r"\d{6}", data.get("pincode", "")):
#         errors.append("Invalid PIN Code")

#     try:
#         datetime.strptime(data.get("dob", ""), "%Y-%m-%d")
#     except:
#         errors.append("Invalid Date of Birth")

#     try:
#         income = int(data.get("annualIncome", 0))
#         if income <= 0:
#             errors.append("Invalid Income Value")
#     except:
#         errors.append("Invalid Income Value")

#     if errors:
#         return jsonify({
#             "data": {
#                 "validation_errors": errors
#             }
#         }), 400

#     # -------- Generate KYC Reference --------
#     kyc_reference = "KYC-" + uuid.uuid4().hex[:10].upper()

#     # -------- Structured RBI-Style Output --------

#     response_data = {
#         "kycReferenceNumber": kyc_reference,
#         "customerIdentification": {
#             "customerId": data.get("customerId"),
#             "fullName": data.get("fullName"),
#             "guardianName": data.get("guardianName"),
#             "dob": data.get("dob"),
#             "gender": data.get("gender"),
#             "nationality": data.get("nationality")
#         },
#         "officiallyValidDocuments": {
#             "maskedPAN": mask_pan(data.get("panNo")),
#             "maskedAadhaar": mask_aadhaar(data.get("aadhaarNo"))
#         },
#         "contactDetails": {
#             "maskedMobile": mask_mobile(data.get("mobile")),
#             "email": data.get("email"),
#             "city": data.get("city"),
#             "state": data.get("state")
#         },
#         "financialProfile": {
#             "occupation": data.get("occupation"),
#             "employer": data.get("employer"),
#             "annualIncome": income,
#             "sourceOfFunds": data.get("sourceOfFunds")
#         },
#         "regulatoryDeclaration": {
#             "pepStatus": data.get("pepStatus"),
#             "taxCountry": data.get("taxCountry")
#         }
#     }

#     return jsonify({
#         "data": response_data
#     })
import os
import json
import pandas as pd
from flask import Flask, request, jsonify, render_template_string, redirect, url_for
from docx import Document
from PyPDF2 import PdfReader

app = Flask(__name__)

# =============================
# CONFIG
# =============================

BASE_INTERNAL = "internal_files"
os.makedirs(BASE_INTERNAL, exist_ok=True)

API_KEY = os.environ.get("API_KEY", "Aravind-sai-0119")

temporary_data = {}

# =============================
# API KEY VALIDATION
# =============================

def require_api_key():
    key = request.headers.get("x-api-key")
    if key != API_KEY:
        return False
    return True

# =============================
# FILE CONVERTER
# =============================

def convert_file(file, filename):
    name = filename.lower()

    try:
        if name.endswith(".csv"):
            df = pd.read_csv(file)
            return df.to_dict(orient="records")

        elif name.endswith(".json"):
            return json.load(file)

        elif name.endswith(".xlsx"):
            sheets = pd.read_excel(file, sheet_name=None)
            return {s: df.to_dict(orient="records") for s, df in sheets.items()}

        elif name.endswith(".txt"):
            return {"text_content": file.read().decode("utf-8", errors="ignore")}

        elif name.endswith(".pdf"):
            reader = PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return {"pdf_text": text}

        elif name.endswith(".docx"):
            doc = Document(file)
            return {"docx_text": "\n".join([p.text for p in doc.paragraphs])}

        else:
            return {"error": "Unsupported file type"}

    except Exception as e:
        return {"error": str(e)}

# =============================
# UI
# =============================

@app.route("/")
def home():
    folders = os.listdir(BASE_INTERNAL)

    return render_template_string("""
<!DOCTYPE html>
<html>
<head>
<title>Enterprise File Manager</title>
<link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css" rel="stylesheet">
<style>
body { background:#f4f6f9; }
.section { background:white; padding:20px; border-radius:10px; margin-bottom:20px; }
</style>
</head>
<body>

<div class="container mt-5">
<h2 class="text-center mb-4">📂 Enterprise File Manager</h2>

<div class="section">
<h5>Temporary Upload (Replace Mode)</h5>
<form action="/upload_temp" method="POST" enctype="multipart/form-data">
<input type="file" name="files" multiple class="form-control mb-2" required>
<button class="btn btn-primary">Upload Temporary</button>
</form>
</div>

<div class="section">
<h5>Upload to Root Internal Folder</h5>
<form action="/upload_internal" method="POST" enctype="multipart/form-data">
<input type="file" name="files" multiple class="form-control mb-2" required>
<button class="btn btn-success">Upload Internal</button>
</form>
</div>

<div class="section">
<h5>Create / Upload to Custom Folder</h5>
<form action="/upload_custom_internal" method="POST" enctype="multipart/form-data">
<input type="text" name="folder_name" placeholder="Folder Name" class="form-control mb-2" required>
<input type="file" name="files" multiple class="form-control mb-2" required>
<button class="btn btn-dark">Upload to Folder</button>
</form>
</div>

<hr>
<h5>Internal Folders</h5>
<ul class="list-group">
{% for folder in folders %}
<li class="list-group-item d-flex justify-content-between">
{{folder}}
<div>
<a href="/api/internal/{{folder}}" class="btn btn-sm btn-outline-primary">View</a>
<a href="/api/internal/delete-folder/{{folder}}" class="btn btn-sm btn-outline-danger">Delete</a>
</div>
</li>
{% endfor %}
</ul>

<hr>
<a href="/view_temp" class="btn btn-dark mt-3">View Temporary JSON</a>

</div>
</body>
</html>
""", folders=folders)

# =============================
# TEMPORARY UPLOAD
# =============================

@app.route("/upload_temp", methods=["POST"])
def upload_temp():
    global temporary_data
    temporary_data = {}

    files = request.files.getlist("files")
    for file in files:
        temporary_data[file.filename] = convert_file(file, file.filename)

    return redirect(url_for("home"))

# =============================
# INTERNAL ROOT UPLOAD
# =============================

@app.route("/upload_internal", methods=["POST"])
def upload_internal():
    files = request.files.getlist("files")

    for file in files:
        path = os.path.join(BASE_INTERNAL, file.filename + ".json")

        if os.path.exists(path):
            continue

        data = convert_file(file, file.filename)
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=4)

    return redirect(url_for("home"))

# =============================
# CUSTOM FOLDER UPLOAD
# =============================

@app.route("/upload_custom_internal", methods=["POST"])
def upload_custom_internal():
    folder_name = request.form.get("folder_name")
    folder_path = os.path.join(BASE_INTERNAL, folder_name)
    os.makedirs(folder_path, exist_ok=True)

    files = request.files.getlist("files")

    for file in files:
        safe_name = os.path.basename(file.filename)
        path = os.path.join(folder_path, safe_name + ".json")

        if os.path.exists(path):
            continue

        data = convert_file(file, safe_name)
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=4)

    return redirect(url_for("home"))

# =============================
# API: LIST ALL INTERNAL
# =============================

@app.route("/api/internal")
def list_internal():
    if not require_api_key():
        return jsonify({"error": "Unauthorized"}), 401

    result = {}
    for folder in os.listdir(BASE_INTERNAL):
        folder_path = os.path.join(BASE_INTERNAL, folder)
        if os.path.isdir(folder_path):
            result[folder] = os.listdir(folder_path)

    return jsonify(result)

# =============================
# API: VIEW FOLDER DATA
# =============================

@app.route("/api/internal/<folder>")
def view_folder(folder):
    if not require_api_key():
        return jsonify({"error": "Unauthorized"}), 401

    folder_path = os.path.join(BASE_INTERNAL, folder)

    if not os.path.exists(folder_path):
        return jsonify({"error": "Folder not found"})

    data = {}

    for file in os.listdir(folder_path):
        if file.endswith(".json"):
            with open(os.path.join(folder_path, file), "r", encoding="utf-8") as f:
                data[file] = json.load(f)

    return jsonify({
        "folder": folder,
        "total_files": len(data),
        "data": data
    })

# =============================
# API: DELETE FILE
# =============================

@app.route("/api/internal/delete-file/<folder>/<file>")
def delete_file(folder, file):
    if not require_api_key():
        return jsonify({"error": "Unauthorized"}), 401

    path = os.path.join(BASE_INTERNAL, folder, file)

    if os.path.exists(path):
        os.remove(path)
        return jsonify({"status": "File deleted"})
    else:
        return jsonify({"error": "File not found"})

# =============================
# API: DELETE FOLDER
# =============================

@app.route("/api/internal/delete-folder/<folder>")
def delete_folder(folder):
    if not require_api_key():
        return jsonify({"error": "Unauthorized"}), 401

    folder_path = os.path.join(BASE_INTERNAL, folder)

    if not os.path.exists(folder_path):
        return jsonify({"error": "Folder not found"})

    for f in os.listdir(folder_path):
        os.remove(os.path.join(folder_path, f))
    os.rmdir(folder_path)

    return jsonify({"status": "Folder deleted"})

# =============================
# API: DELETE ALL
# =============================

@app.route("/api/internal/delete-all")
def delete_all():
    if not require_api_key():
        return jsonify({"error": "Unauthorized"}), 401

    for folder in os.listdir(BASE_INTERNAL):
        folder_path = os.path.join(BASE_INTERNAL, folder)
        if os.path.isdir(folder_path):
            for f in os.listdir(folder_path):
                os.remove(os.path.join(folder_path, f))
            os.rmdir(folder_path)

    return jsonify({"status": "All internal data deleted"})

# =============================
# VIEW TEMP
# =============================

@app.route("/view_temp")
def view_temp():
    if not temporary_data:
        return jsonify({"message": "No temporary files uploaded"})
    return jsonify(temporary_data)

# =============================
# RUN
# =============================

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)

